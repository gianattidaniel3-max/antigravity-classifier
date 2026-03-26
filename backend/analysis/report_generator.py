"""
Report generator for Analysis Runs.
Supports PDF (via reportlab) and DOCX (via python-docx).
"""
from __future__ import annotations
import io
import datetime
from typing import Dict, List, Any, Optional


# ── Shared helpers ────────────────────────────────────────────────────────────

def _format_run_date(run_at: str) -> str:
    try:
        return datetime.datetime.fromisoformat(run_at).strftime("%d/%m/%Y %H:%M")
    except Exception:
        return run_at


def _format_flags(triggered_rules: List[Dict]) -> List[str]:
    return [r.get("flag_label") or f"{r.get('field')} {r.get('op')}" for r in triggered_rules]


def _extract_run_data(run: Dict):
    """Return (summary, results, formatted_date) from a run dict."""
    return (
        run.get("summary", {}),
        run.get("results", []),
        _format_run_date(run.get("run_at", "")),
    )


NO_RESULTS_MESSAGE = "Nessun documento ha soddisfatto le regole."


# ── PDF ──────────────────────────────────────────────────────────────────────

def generate_pdf(run: Dict, case_name: str, template_name: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title", parent=styles["Title"],
        fontSize=18, textColor=colors.HexColor("#18181b"), spaceAfter=4,
    )
    sub_style = ParagraphStyle(
        "Sub", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#71717a"), spaceAfter=2,
    )
    heading_style = ParagraphStyle(
        "Heading", parent=styles["Heading2"],
        fontSize=11, textColor=colors.HexColor("#4f46e5"), spaceBefore=12, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#334155"),
    )
    flag_style = ParagraphStyle(
        "Flag", parent=styles["Normal"],
        fontSize=8, textColor=colors.HexColor("#92400e"),
    )

    summary, results, run_at = _extract_run_data(run)
    story: List[Any] = []

    # Header
    story.append(Paragraph("ECO - Extractor", title_style))
    story.append(Paragraph("Report Analisi Fascicolo", sub_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceAfter=8))

    # Meta table
    meta = [
        ["Fascicolo:", case_name],
        ["Template:", template_name],
        ["Data analisi:", run_at],
        ["Documenti totali:", str(summary.get("total", 0))],
        ["Documenti segnalati:", str(summary.get("flagged", 0))],
    ]
    meta_table = Table(meta, colWidths=[4*cm, 12*cm])
    meta_table.setStyle(TableStyle([
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("TEXTCOLOR", (0,0), (0,-1), colors.HexColor("#71717a")),
        ("TEXTCOLOR", (1,0), (1,-1), colors.HexColor("#18181b")),
        ("FONTNAME", (0,0), (0,-1), "Helvetica-Bold"),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ("TOPPADDING", (0,0), (-1,-1), 3),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 0.5*cm))

    # Global Insights
    global_insights = run.get("global_insights")
    if global_insights and isinstance(global_insights, list):
        for item in global_insights:
            name = item.get("name", "AI Insight Globale")
            insight = item.get("insight")
            if not insight: continue
            
            story.append(Paragraph(name, heading_style))
            insight_p = Paragraph(insight.replace("\n", "<br/>"), body_style)
            insight_table = Table([[insight_p]], colWidths=[16*cm])
            insight_table.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#f0fdf4")),
                ("BOX", (0,0), (-1,-1), 0.5, colors.HexColor("#a7d3bc")),
                ("TOPPADDING", (0,0), (-1,-1), 8),
                ("BOTTOMPADDING", (0,0), (-1,-1), 8),
                ("LEFTPADDING", (0,0), (-1,-1), 8),
                ("RIGHTPADDING", (0,0), (-1,-1), 8),
            ]))
            story.append(insight_table)
            story.append(Spacer(1, 0.5*cm))

    # By-rule summary
    by_rule = summary.get("by_rule", {})
    if by_rule:
        story.append(Paragraph("Sommario per regola", heading_style))
        rule_data = [["Regola", "Documenti"]] + [[k, str(v)] for k, v in by_rule.items()]
        rule_table = Table(rule_data, colWidths=[12*cm, 4*cm])
        rule_table.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#4f46e5")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#f8fafc"), colors.white]),
            ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("ALIGN", (1,0), (1,-1), "CENTER"),
        ]))
        story.append(rule_table)
        story.append(Spacer(1, 0.4*cm))

    # Flagged documents
    if results:
        story.append(Paragraph("Documenti segnalati", heading_style))
        for row in results:
            flags = _format_flags(row.get("triggered_rules", []))
            doc_data: List[Any] = [[Paragraph(f"<b>{row.get('filename', '—')}</b>", body_style), ""]]
            label, date = row.get("label", ""), row.get("date", "")
            if label or date:
                doc_data.append([Paragraph(label, body_style), Paragraph(date, body_style)])
            if flags:
                doc_data.append([Paragraph("Flags: " + " · ".join(flags), flag_style), ""])

            doc_table = Table(doc_data, colWidths=[10*cm, 6*cm])
            doc_table.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#fffbeb")),
                ("BOX", (0,0), (-1,-1), 0.5, colors.HexColor("#fde68a")),
                ("FONTSIZE", (0,0), (-1,-1), 9),
                ("BOTTOMPADDING", (0,0), (-1,-1), 3),
                ("TOPPADDING", (0,0), (-1,-1), 3),
                ("LEFTPADDING", (0,0), (-1,-1), 6),
            ]))
            story.append(doc_table)
            story.append(Spacer(1, 0.2*cm))
    else:
        story.append(Paragraph(NO_RESULTS_MESSAGE, body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


# ── DOCX ─────────────────────────────────────────────────────────────────────

def generate_docx(run: Dict, case_name: str, template_name: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    summary, results, run_at = _extract_run_data(run)

    # Title
    title = document.add_heading("ECO - Extractor", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_p = document.add_paragraph("Report Analisi Fascicolo")
    run_p.runs[0].font.color.rgb = RGBColor(0x71, 0x71, 0x7a)
    run_p.runs[0].font.size = Pt(10)
    document.add_paragraph()

    # Meta
    for label, value in [
        ("Fascicolo", case_name),
        ("Template", template_name),
        ("Data analisi", run_at),
        ("Documenti totali", str(summary.get("total", 0))),
        ("Documenti segnalati", str(summary.get("flagged", 0))),
    ]:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_label = p.add_run(f"{label}: ")
        r_label.bold = True
        r_label.font.size = Pt(10)
        r_val = p.add_run(value)
        r_val.font.size = Pt(10)

    # Global Insights
    global_insights = run.get("global_insights")
    if global_insights and isinstance(global_insights, list):
        for item in global_insights:
            name = item.get("name", "AI Insight Globale")
            insight = item.get("insight")
            if not insight: continue
            
            document.add_heading(name, level=2)
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run(insight)
            r.font.size = Pt(10)
            r.bold = False 

    document.add_paragraph()

    # By-rule summary
    by_rule = summary.get("by_rule", {})
    if by_rule:
        document.add_heading("Sommario per regola", level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Regola"
        hdr[1].text = "Documenti"
        for cell in hdr:
            for r in cell.paragraphs[0].runs:
                r.bold = True
        for k, v in by_rule.items():
            row = table.add_row().cells
            row[0].text = k
            row[1].text = str(v)
        document.add_paragraph()

    # Flagged documents
    if results:
        document.add_heading("Documenti segnalati", level=2)
        for row in results:
            flags = _format_flags(row.get("triggered_rules", []))
            label, date = row.get("label", ""), row.get("date", "")

            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r_fn = p.add_run(row.get("filename", "—"))
            r_fn.bold = True
            r_fn.font.size = Pt(10)

            if label or date:
                p2 = document.add_paragraph("  ".join(x for x in [label, date] if x))
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(1)
                p2.runs[0].font.size = Pt(9)
                p2.runs[0].font.color.rgb = RGBColor(0x71, 0x71, 0x7a)

            if flags:
                p3 = document.add_paragraph("Flags: " + " · ".join(flags))
                p3.paragraph_format.space_after = Pt(6)
                p3.runs[0].font.size = Pt(9)
                p3.runs[0].font.color.rgb = RGBColor(0x92, 0x40, 0x0e)
    else:
        document.add_paragraph(NO_RESULTS_MESSAGE)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()
