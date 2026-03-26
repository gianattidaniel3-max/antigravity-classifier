from fastapi import APIRouter, Depends, HTTPException
from backend.auth.deps import get_current_user
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from typing import List
import uuid, io, csv, json, datetime

import requests as http_requests

from backend.db.session import get_db
from backend.db.models import Case, Document, AnalysisTemplate, AnalysisRun, CaseStatus, User
from backend.analysis.rule_engine import run_analysis, format_case_context
from backend.analysis.report_generator import generate_pdf, generate_docx

_OLLAMA_URL  = "http://localhost:11434/api/chat"
_OLLAMA_MODEL = "llama3.2:3b"
_NLP_FIELDS = (
    "label (tipo documento), category (categoria), date (data del documento), "
    "importo (importo in euro), mittente (chi ha inviato), destinatario (chi ha ricevuto), "
    "oggetto (argomento), scadenza (data di scadenza), tribunale (tribunale competente), "
    "numero_decreto (numero decreto), numero_rg (numero R.G.)"
)
_NLP_OPS = (
    "eq (uguale a), neq (diverso da), contains (contiene nel testo), "
    "gt (maggiore di — per numeri e date), lt (minore di — per numeri e date), "
    "is_null (campo assente/vuoto), not_null (campo presente), "
    "in (uno tra una lista, valori separati da virgola)"
)


router = APIRouter(dependencies=[Depends(get_current_user)])


# ─────────────────────────────────────────────────────────────────
# Helper
# ─────────────────────────────────────────────────────────────────

def _doc_to_dict(doc: Document) -> dict:
    return {
        "id":                doc.id,
        "filename":          doc.filename,
        "extracted_label":   doc.extracted_label,
        "extracted_category": doc.extracted_category,
        "extracted_date":    doc.extracted_date,
        "extracted_fields":  doc.extracted_fields or {},
        "confidence_score":  doc.confidence_score,
        "status":            doc.status.value if doc.status else "pending",
        "human_verified":    doc.human_verified,
        "llm_notes":         doc.llm_notes,
        "llm_classification_match": doc.llm_classification_match,
    }


def _case_to_dict(case: Case, include_docs: bool = False) -> dict:
    d = {
        "id":          case.id,
        "name":        case.name,
        "description": case.description,
        "client_name": case.client_name,
        "status":      case.status.value,
        "created_at":  case.created_at.isoformat() if case.created_at else None,
        "doc_count":   len(case.documents),
    }
    if include_docs:
        d["documents"] = [_doc_to_dict(doc) for doc in case.documents]
    return d


# ─────────────────────────────────────────────────────────────────
# Cases CRUD
# ─────────────────────────────────────────────────────────────────

@router.get("/cases")
def list_cases(db: Session = Depends(get_db)):
    cases = db.query(Case).order_by(Case.created_at.desc()).all()
    return [_case_to_dict(c) for c in cases]


@router.post("/cases")
def create_case(payload: dict, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    name = payload.get("name", "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="name required")
    case = Case(
        id=str(uuid.uuid4()),
        name=name,
        description=payload.get("description"),
        client_name=payload.get("client_name"),
        created_by=current_user.id
    )
    db.add(case)
    db.commit()
    db.refresh(case)
    return _case_to_dict(case)


@router.get("/cases/{case_id}")
def get_case(case_id: str, db: Session = Depends(get_db)):
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    return _case_to_dict(case, include_docs=True)


@router.patch("/cases/{case_id}")
def update_case(case_id: str, payload: dict, db: Session = Depends(get_db)):
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    for field in ("name", "description", "client_name"):
        if field in payload:
            setattr(case, field, payload[field])
    if "status" in payload:
        try:
            case.status = CaseStatus(payload["status"])
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid status: {payload['status']}")
    db.commit()
    return _case_to_dict(case)


@router.delete("/cases/{case_id}")
def delete_case(case_id: str, db: Session = Depends(get_db)):
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    # Detach documents rather than deleting them
    for doc in case.documents:
        doc.case_id = None
    db.delete(case)
    db.commit()
    return {"deleted": case_id}


# ─────────────────────────────────────────────────────────────────
# Document ↔ Case assignment
# ─────────────────────────────────────────────────────────────────

@router.post("/cases/{case_id}/documents/{doc_id}")
def assign_document(case_id: str, doc_id: str, db: Session = Depends(get_db)):
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    doc.case_id = case_id
    db.commit()
    return {"assigned": doc_id, "case": case_id}


@router.delete("/cases/{case_id}/documents/{doc_id}")
def unassign_document(case_id: str, doc_id: str, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id, Document.case_id == case_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found in this case")
    doc.case_id = None
    db.commit()
    return {"unassigned": doc_id}


# ─────────────────────────────────────────────────────────────────
# Unassigned documents (for the picker)
# ─────────────────────────────────────────────────────────────────

@router.get("/documents-unassigned")
def list_unassigned(db: Session = Depends(get_db)):
    docs = db.query(Document).filter(Document.case_id == None).order_by(Document.upload_date.desc()).all()
    return [_doc_to_dict(d) for d in docs]


# ─────────────────────────────────────────────────────────────────
# Analysis Templates CRUD
# ─────────────────────────────────────────────────────────────────

@router.get("/templates")
def list_templates(db: Session = Depends(get_db)):
    return [
        {
            "id":          t.id,
            "name":        t.name,
            "description": t.description,
            "rules":       t.rules,
            "global_prompts": t.global_prompts,
            "created_at":  t.created_at.isoformat() if t.created_at else None,
        }
        for t in db.query(AnalysisTemplate).order_by(AnalysisTemplate.created_at).all()
    ]


@router.post("/templates")
def create_template(payload: dict, db: Session = Depends(get_db)):
    name = payload.get("name", "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="name required")
    if db.query(AnalysisTemplate).filter(AnalysisTemplate.name == name).first():
        raise HTTPException(status_code=409, detail=f"Template '{name}' already exists")
    tmpl = AnalysisTemplate(
        id=str(uuid.uuid4()),
        name=name,
        description=payload.get("description"),
        rules=payload.get("rules", []),
    )
    db.add(tmpl)
    db.commit()
    db.refresh(tmpl)
    return {"id": tmpl.id, "name": tmpl.name, "rules": tmpl.rules, "global_prompts": tmpl.global_prompts}


@router.put("/templates/{tmpl_id}")
def update_template(tmpl_id: str, payload: dict, db: Session = Depends(get_db)):
    tmpl = db.query(AnalysisTemplate).filter(AnalysisTemplate.id == tmpl_id).first()
    if not tmpl:
        raise HTTPException(status_code=404, detail="Template not found")
    if "name" in payload:
        tmpl.name = payload["name"]
    if "description" in payload:
        tmpl.description = payload["description"]
    if "rules" in payload:
        tmpl.rules = payload["rules"]
    if "global_prompts" in payload:
        tmpl.global_prompts = payload["global_prompts"]
    db.commit()
    return {"id": tmpl.id, "name": tmpl.name, "rules": tmpl.rules, "global_prompts": tmpl.global_prompts}


@router.delete("/templates/{tmpl_id}")
def delete_template(tmpl_id: str, db: Session = Depends(get_db)):
    tmpl = db.query(AnalysisTemplate).filter(AnalysisTemplate.id == tmpl_id).first()
    if not tmpl:
        raise HTTPException(status_code=404, detail="Template not found")
    db.delete(tmpl)
    db.commit()
    return {"deleted": tmpl_id}


# ─────────────────────────────────────────────────────────────────
# Run analysis
# ─────────────────────────────────────────────────────────────────

@router.post("/cases/{case_id}/run/{tmpl_id}")
def run_case_analysis(case_id: str, tmpl_id: str, db: Session = Depends(get_db)):
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    tmpl = db.query(AnalysisTemplate).filter(AnalysisTemplate.id == tmpl_id).first()
    if not tmpl:
        raise HTTPException(status_code=404, detail="Template not found")

    docs = [_doc_to_dict(d) for d in case.documents]
    results, summary = run_analysis(docs, tmpl.rules)
    
    global_insights = []
    if tmpl.global_prompts and isinstance(tmpl.global_prompts, list):
        context = format_case_context(docs)
        for item in tmpl.global_prompts:
            name = item.get("name", "Analisi")
            prompt_text = item.get("prompt")
            if not prompt_text: continue

            prompt = (
                "Sei un assistente legale esperto. Analizza i dati del fascicolo forniti "
                "e rispondi alla domanda dell'utente in modo professionale e sintetico.\n\n"
                f"{context}\n\n"
                f"Domanda Analitica: {prompt_text}"
            )
            try:
                resp = http_requests.post(
                    _OLLAMA_URL,
                    json={
                        "model": _OLLAMA_MODEL,
                        "messages": [{"role": "user", "content": prompt}],
                        "stream": False,
                        "options": {"temperature": 0.1},
                    },
                    timeout=60,
                )
                resp.raise_for_status()
                insight = resp.json()["message"]["content"]
                global_insights.append({"name": name, "insight": insight})
            except Exception as e:
                global_insights.append({"name": name, "insight": f"Errore: {e}"})

    run = AnalysisRun(
        id=str(uuid.uuid4()),
        case_id=case_id,
        template_id=tmpl_id,
        results=results,
        summary=summary,
        global_insights=global_insights,
    )
    db.add(run)
    db.commit()

    return {
        "run_id":   run.id,
        "run_at":   run.run_at.isoformat(),
        "summary":  summary,
        "results":  results,
        "global_insights": global_insights,
    }


@router.get("/cases/{case_id}/runs")
def list_runs(case_id: str, db: Session = Depends(get_db)):
    runs = (
        db.query(AnalysisRun)
        .filter(AnalysisRun.case_id == case_id)
        .order_by(AnalysisRun.run_at.desc())
        .all()
    )
    return [
        {
            "id":          r.id,
            "template_id": r.template_id,
            "run_at":      r.run_at.isoformat(),
            "summary":     r.summary,
            "global_insights": r.global_insights,
        }
        for r in runs
    ]


@router.get("/runs/{run_id}")
def get_run(run_id: str, db: Session = Depends(get_db)):
    run = db.query(AnalysisRun).filter(AnalysisRun.id == run_id).first()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    return {
        "id":          run.id,
        "case_id":     run.case_id,
        "template_id": run.template_id,
        "run_at":      run.run_at.isoformat(),
        "summary":     run.summary,
        "results":     run.results,
        "global_insights": run.global_insights,
    }


# ─────────────────────────────────────────────────────────────────
# Export helpers
# ─────────────────────────────────────────────────────────────────

def _get_run_or_404(run_id: str, db: Session) -> AnalysisRun:
    run = db.query(AnalysisRun).filter(AnalysisRun.id == run_id).first()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    return run


def _run_to_export_dict(run: AnalysisRun) -> dict:
    return {"run_at": run.run_at.isoformat(), "summary": run.summary, "results": run.results}


def _file_response(data: bytes, media_type: str, filename: str) -> StreamingResponse:
    return StreamingResponse(
        io.BytesIO(data),
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ─────────────────────────────────────────────────────────────────
# Case dataset exports  (raw data — not tied to a specific run)
# ─────────────────────────────────────────────────────────────────

_DATASET_COLS = [
    ("Filename",     None,                  "filename"),
    ("Categoria",    None,                  "extracted_category"),
    ("Tipo",         None,                  "extracted_label"),
    ("Data",         None,                  "extracted_date"),
    ("Importo",      "importo",             None),
    ("Mittente",     "mittente",            None),
    ("Destinatario", "destinatario",        None),
    ("Oggetto",      "oggetto",             None),
    ("Scadenza",     "scadenza",            None),
    ("Tribunale",    "tribunale",           None),
    ("N. Decreto",   "numero_decreto",      None),
    ("N. R.G.",      "numero_rg",           None),
    ("Confidenza",   None,                  "confidence_score"),
    ("Verificato",   None,                  "human_verified"),
    ("Status",       None,                  "status"),
]

def _doc_row(doc: Document) -> list:
    fields = doc.extracted_fields or {}
    row = []
    for _, field_key, attr_key in _DATASET_COLS:
        if field_key:
            row.append(fields.get(field_key) or "")
        elif attr_key == "confidence_score":
            row.append(f"{doc.confidence_score:.0%}" if doc.confidence_score else "")
        elif attr_key == "human_verified":
            row.append("Sì" if doc.human_verified else "No")
        elif attr_key == "status":
            row.append(doc.status.value if doc.status else "")
        else:
            row.append(getattr(doc, attr_key, "") or "")
    return row


@router.get("/cases/{case_id}/export-dataset-excel")
def export_case_dataset_excel(case_id: str, db: Session = Depends(get_db)):
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Dataset"

    header_fill = PatternFill("solid", fgColor="2d6a4f")
    header_font = Font(bold=True, color="FFFFFF", size=10)
    even_fill   = PatternFill("solid", fgColor="F8FAFC")

    headers = [c[0] for c in _DATASET_COLS]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    for row_idx, doc in enumerate(case.documents, 2):
        for col, value in enumerate(_doc_row(doc), 1):
            cell = ws.cell(row=row_idx, column=col, value=value)
            if row_idx % 2 == 0:
                cell.fill = even_fill
            cell.alignment = Alignment(vertical="center", wrap_text=True)

    for col in range(1, len(headers) + 1):
        max_len = max(
            len(str(ws.cell(row=r, column=col).value or ""))
            for r in range(1, len(case.documents) + 2)
        )
        ws.column_dimensions[get_column_letter(col)].width = min(max_len + 4, 42)

    ws.freeze_panes = "A2"
    ws.row_dimensions[1].height = 18

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    safe_name = "".join(c for c in case.name if c.isalnum() or c in (" ", "_", "-")).replace(" ", "_").strip()[:50]
    return _file_response(
        buf.read(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        f"{safe_name}.xlsx",
    )


@router.get("/cases/{case_id}/export-dataset-docx")
def export_case_dataset_docx(case_id: str, db: Session = Depends(get_db)):
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    document = DocxDoc()
    for section in document.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        # Landscape A4
        section.orientation = 1
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21)

    title = document.add_heading(f"Dataset – {case.name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = document.add_paragraph(
        f"Cliente: {case.client_name or '—'}   |   Documenti: {len(case.documents)}"
    )
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x71, 0x71, 0x7A)
    document.add_paragraph()

    # Narrower set of columns for readability in landscape A4
    DOC_COLS = [c for c in _DATASET_COLS if c[0] not in ("Oggetto", "Status")]
    headers = [c[0] for c in DOC_COLS]

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = col_name
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "2d6a4f")
        shd.set(qn("w:val"),  "clear")
        tcPr.append(shd)

    for doc in case.documents:
        full_row = _doc_row(doc)
        # Map to narrower column set
        col_indices = [i for i, c in enumerate(_DATASET_COLS) if c[0] in headers]
        row_cells = table.add_row().cells
        for j, idx in enumerate(col_indices):
            row_cells[j].text = str(full_row[idx])
            for para in row_cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7.5)

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    safe_name = "".join(c for c in case.name if c.isalnum() or c in (" ", "_", "-")).replace(" ", "_").strip()[:50]
    return _file_response(
        buf.read(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{safe_name}.docx",
    )


# ─────────────────────────────────────────────────────────────────
# CSV export
# ─────────────────────────────────────────────────────────────────

@router.get("/runs/{run_id}/export-csv")
def export_run_csv(run_id: str, db: Session = Depends(get_db)):
    run = _get_run_or_404(run_id, db)

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["document_id", "filename", "label", "category", "date", "triggered_flags"])

    for row in run.results:
        flags = "; ".join(r.get("flag_label", "") for r in row.get("triggered_rules", []))
        writer.writerow([
            row.get("document_id", ""),
            row.get("filename", ""),
            row.get("label", ""),
            row.get("category", ""),
            row.get("date", ""),
            flags,
        ])

    csv_bytes = output.getvalue().encode("utf-8")
    case = db.query(Case).filter(Case.id == run.case_id).first()
    safe_name = "".join(c for c in case.name if c.isalnum() or c in (" ", "_", "-")).replace(" ", "_").strip()[:50] if case else f"analisi_{run_id[:8]}"
    return _file_response(csv_bytes, "text/csv", f"{safe_name}.csv")


@router.get("/runs/{run_id}/export-pdf")
def export_run_pdf(run_id: str, db: Session = Depends(get_db)):
    run = _get_run_or_404(run_id, db)
    case = db.query(Case).filter(Case.id == run.case_id).first()
    template = db.query(AnalysisTemplate).filter(AnalysisTemplate.id == run.template_id).first()
    pdf_bytes = generate_pdf(
        _run_to_export_dict(run),
        case.name if case else "—",
        template.name if template else "—",
    )
    safe_name = "".join(c for c in case.name if c.isalnum() or c in (" ", "_", "-")).replace(" ", "_").strip()[:50] if case else f"analisi_{run_id[:8]}"
    return _file_response(pdf_bytes, "application/pdf", f"{safe_name}.pdf")


@router.get("/runs/{run_id}/export-docx")
def export_run_docx(run_id: str, db: Session = Depends(get_db)):
    run = _get_run_or_404(run_id, db)
    case = db.query(Case).filter(Case.id == run.case_id).first()
    template = db.query(AnalysisTemplate).filter(AnalysisTemplate.id == run.template_id).first()
    docx_bytes = generate_docx(
        _run_to_export_dict(run),
        case.name if case else "—",
        template.name if template else "—",
    )
    safe_name = "".join(c for c in case.name if c.isalnum() or c in (" ", "_", "-")).replace(" ", "_").strip()[:50] if case else f"analisi_{run_id[:8]}"
    return _file_response(
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{safe_name}.docx",
    )


# ─────────────────────────────────────────────────────────────────
# NLP → rule (Llama)
# ─────────────────────────────────────────────────────────────────

@router.post("/nlp-to-rule")
def nlp_to_rule(payload: dict):
    """Convert natural language description to a structured rule via Llama."""
    text = (payload.get("text") or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="text required")

    prompt = (
        "Sei un assistente per uno studio legale italiano che analizza documenti legali. "
        "L'utente vuole creare una regola per identificare documenti che soddisfano un certo criterio.\n\n"
        f"Campi disponibili: {_NLP_FIELDS}\n"
        f"Operatori disponibili: {_NLP_OPS}\n\n"
        f"Descrizione dell'utente: \"{text}\"\n\n"
        "Restituisci SOLO un oggetto JSON con questi campi:\n"
        "- field: il nome del campo (stringa)\n"
        "- op: l'operatore (stringa)\n"
        "- value: il valore da confrontare (stringa o null se non applicabile)\n"
        "- flag_label: etichetta breve in italiano che descrive cosa segnala questa regola (max 4 parole)\n\n"
        "Esempio di output: {\"field\": \"importo\", \"op\": \"gt\", \"value\": \"10000\", \"flag_label\": \"Importo elevato\"}"
    )

    try:
        resp = http_requests.post(
            _OLLAMA_URL,
            json={
                "model": _OLLAMA_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "format": "json",
                "stream": False,
                "options": {"temperature": 0.05, "num_predict": 128},
            },
            timeout=45,
        )
        resp.raise_for_status()
        rule = json.loads(resp.json()["message"]["content"])
        if "field" not in rule or "op" not in rule:
            raise ValueError("incomplete rule")
        return {
            "field":      str(rule.get("field", "label")),
            "op":         str(rule.get("op", "contains")),
            "value":      str(rule["value"]).strip() if rule.get("value") not in (None, "null", "") else "",
            "flag_label": str(rule.get("flag_label", text[:30])),
        }
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"AI non disponibile: {e}")
